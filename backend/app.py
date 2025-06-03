from dotenv import load_dotenv
import os

# Force reload of .env file
load_dotenv(override=True)

# Debug: Print environment variables (without showing full key)
print("🔑 Checking API key configuration...")
openai_api_key = os.environ.get("OPENAI_API_KEY")
google_api_key = os.environ.get("GOOGLE_API_KEY")

if openai_api_key:
    if openai_api_key.startswith("sk-"):
        print(f"✅ OpenAI API key found (length: {len(openai_api_key)})")
        # Show first 4 and last 4 characters of the key
        masked_key = f"{openai_api_key[:4]}...{openai_api_key[-4:]}" if len(openai_api_key) > 8 else "***"
        print(f"🔐 OpenAI API Key: {masked_key}")
    else:
        print("❌ Invalid OpenAI API key format! Key should start with 'sk-'")
        print(f"Current key starts with: {openai_api_key[:4]}")
else:
    print("❌ No OpenAI API key found in environment variables!")
    print("Please check your .env file and ensure it contains OPENAI_API_KEY")

if google_api_key:
    print(f"✅ Google API key found (length: {len(google_api_key)})")
    masked_key = f"{google_api_key[:4]}...{google_api_key[-4:]}" if len(google_api_key) > 8 else "***"
    print(f"🔐 Google API Key: {masked_key}")
else:
    print("❌ No Google API key found in environment variables!")

import openai
from google import generativeai as genai

# Configure Google API
genai.configure(api_key=google_api_key)

from fastapi import FastAPI, File, UploadFile, Form, Depends, HTTPException, status
from fastapi.responses import JSONResponse, StreamingResponse
from fastapi.middleware.cors import CORSMiddleware
from typing import List, Optional
import uvicorn
import tempfile
import os
import subprocess
import time
import random
import math
import concurrent.futures
import base64
import re
import requests
import json
import io
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import threading

app = FastAPI(title="PV Generation API")

# Configuration CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000"],  # URL de votre frontend Next.js
    allow_credentials=True,
    allow_methods=["*"],  # Autorise toutes les méthodes HTTP
    allow_headers=["*"],  # Autorise tous les headers
)

# --- Helper Functions ---

def extract_file_id_from_url(url):
    patterns = [
        r"https://drive\.google\.com/file/d/([a-zA-Z0-9_-]+)",
        r"https://drive\.google\.com/open\?id=([a-zA-Z0-9_-]+)",
        r"https://drive\.google\.com/uc\?id=([a-zA-Z0-9_-]+)",
        r"id=([a-zA-Z0-9_-]+)"
    ]
    for pattern in patterns:
        match = re.search(pattern, url)
        if match:
            return match.group(1)
    return None

def download_video_from_drive(video_url, output_path):
    try:
        file_id = extract_file_id_from_url(video_url)
        if not file_id:
            return False, "URL Google Drive non reconnue."
        session = requests.Session()
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
            'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8',
            'Accept-Language': 'fr,fr-FR;q=0.8,en-US;q=0.5,en;q=0.3',
            'Accept-Encoding': 'gzip, deflate, br',
            'Connection': 'keep-alive',
            'Upgrade-Insecure-Requests': '1'
        }
        download_url = f'https://drive.usercontent.google.com/download?id={file_id}&export=download&authuser=0&confirm=t'
        response = session.get(download_url, headers=headers, stream=True, timeout=30)
        content_type = response.headers.get('Content-Type', '').lower()
        if 'text/html' in content_type:
            # Try alternative URL for large files
            download_url = f'https://drive.usercontent.google.com/download?id={file_id}&export=download&authuser=0&confirm=t&uuid=123&at=123'
            response = session.get(download_url, headers=headers, stream=True, timeout=30)
            content_type = response.headers.get('Content-Type', '').lower()
            if 'text/html' in content_type:
                return False, "Impossible d'accéder au fichier. Vérifiez les droits de partage."
        temp_path = output_path + ".tmp"
        try:
            chunk_size = 500 * 1024 * 1024
            downloaded_size = 0
            expected_size = None
            if 'content-length' in response.headers:
                expected_size = int(response.headers['content-length'])
            with open(temp_path, 'wb') as f:
                for chunk in response.iter_content(chunk_size=chunk_size):
                    if chunk:
                        f.write(chunk)
                        downloaded_size += len(chunk)
            # Check file
            if os.path.exists(temp_path):
                file_size = os.path.getsize(temp_path)
                if file_size < 10000:
                    os.remove(temp_path)
                    return False, "Fichier téléchargé trop petit."
                # Optionally: check VRO header (not enforced here)
                try:
                    if os.path.exists(output_path):
                        os.remove(output_path)
                    os.rename(temp_path, output_path)
                except Exception as e:
                    import shutil
                    try:
                        shutil.copy2(temp_path, output_path)
                        os.remove(temp_path)
                    except Exception as e2:
                        return False, f"Erreur lors de la copie: {str(e2)}"
                return True, None
            else:
                return False, "Erreur lors de l'écriture du fichier."
        except Exception as e:
            if os.path.exists(temp_path):
                os.remove(temp_path)
            return False, f"Erreur pendant le téléchargement: {str(e)}"
    except Exception as e:
        if os.path.exists(output_path):
            os.remove(output_path)
        return False, f"Erreur inattendue: {str(e)}"

def verify_video_file(file_path):
    """Check if the video file is valid using ffprobe."""
    if not os.path.exists(file_path):
        return False, f"File {file_path} does not exist."
    file_size = os.path.getsize(file_path)
    if file_size < 10000:
        return False, "File is too small to be a valid video."
    probe_command = [
        "ffprobe", "-v", "error", "-show_format", "-show_streams", file_path
    ]
    result = subprocess.run(probe_command, capture_output=True, text=True)
    if result.returncode != 0:
        return False, f"Invalid video format: {result.stderr}"
    return True, None

def extract_audio_from_video(input_video_path, output_audio_path):
    """Extract audio from video using ffmpeg."""
    if not os.path.exists(input_video_path):
        return False, "Video file does not exist."
    if os.path.getsize(input_video_path) == 0:
        return False, "Video file is empty."
    # Convert VRO to MP4 if needed (not implemented here)
    extract_command = [
        'ffmpeg', '-i', input_video_path, '-vn', '-acodec', 'libmp3lame',
        '-ar', '44100', '-ab', '192k', '-y', output_audio_path
    ]
    result = subprocess.run(extract_command, capture_output=True, text=True)
    if result.returncode != 0:
        return False, f"Audio extraction error: {result.stderr}"
    if not os.path.exists(output_audio_path) or os.path.getsize(output_audio_path) == 0:
        return False, "Audio file not created or is empty."
    return True, None

def segment_audio(audio_path, segment_length_ms=120000):
    """Split audio into segments using ffmpeg."""
    try:
        result = subprocess.run([
            'ffprobe', '-v', 'error', '-show_entries',
            'format=duration', '-of',
            'default=noprint_wrappers=1:nokey=1', audio_path
        ], stdout=subprocess.PIPE, stderr=subprocess.STDOUT)
        total_duration = float(result.stdout)
        segment_length_sec = segment_length_ms / 1000
        num_segments = math.ceil(total_duration / segment_length_sec)
        segment_paths = []
        temp_dir = tempfile.gettempdir()
        for i in range(num_segments):
            start_time = i * segment_length_sec
            temp_segment_path = os.path.join(temp_dir, f"segment_{i+1}_{os.path.basename(audio_path)}.mp3")
            extract_cmd = [
                "ffmpeg", "-y", "-i", audio_path, "-ss", str(start_time),
                "-t", str(segment_length_sec), "-c", "copy", temp_segment_path
            ]
            subprocess.run(extract_cmd, stdout=subprocess.DEVNULL, stderr=subprocess.STDOUT)
            if os.path.exists(temp_segment_path):
                segment_paths.append(temp_segment_path)
        return segment_paths
    except Exception as e:
        return []

def transcribe_audio_segments(segments, batch_size=8, timeout=30):
    """Transcribe audio segments using OpenAI's Whisper API with parallel processing."""
    full_transcript = []
    failed_segments = []
    active_threads = 0
    max_active_threads = 0
    
    # Initialize OpenAI client with explicit API key check
    if not openai_api_key or not openai_api_key.startswith("sk-"):
        raise ValueError("Invalid OpenAI API key. Key should start with 'sk-'")
    
    client = openai.OpenAI(api_key=openai_api_key)
    print("✅ OpenAI client initialized successfully")
    
    # Semaphore to limit concurrent API calls
    api_semaphore = threading.Semaphore(5)  # Allow 5 concurrent API calls
    
    def process_segment(segment_info):
        nonlocal active_threads, max_active_threads
        i, segment_path = segment_info
        max_retries = 5
        initial_retry_delay = 5
        retry_delay = initial_retry_delay
        
        active_threads += 1
        max_active_threads = max(max_active_threads, active_threads)
        print(f"🔄 Starting segment {i+1} (Active threads: {active_threads})")
        
        try:
            for attempt in range(max_retries):
                try:
                    with open(segment_path, "rb") as audio_file:
                        print(f"📊 Segment {i+1} size: {os.path.getsize(segment_path)} bytes")
                        
                        with api_semaphore:
                            print(f"🎯 Attempting transcription for segment {i+1} (attempt {attempt + 1}/{max_retries})...")
                            
                            # Use the pre-initialized client
                            response = client.audio.transcriptions.create(
                                model="whisper-1",
                                file=audio_file,
                                language="fr",
                                response_format="text"
                            )
                    
                    if response:
                        print(f"✅ Successfully transcribed segment {i+1}")
                        os.remove(segment_path)
                        return (i, response)
                    else:
                        print(f"⚠️ Segment {i+1} returned no text from Whisper (attempt {attempt + 1})")
                        
                except Exception as e:
                    error_msg = str(e)
                    print(f"❌ Error transcribing segment {i+1} (attempt {attempt + 1}): {error_msg}")
                    
                    if "rate_limit_exceeded" in error_msg.lower():
                        try:
                            import re
                            retry_after = re.search(r'retry after (\d+)', error_msg.lower())
                            if retry_after:
                                retry_delay = int(retry_after.group(1))
                            else:
                                retry_delay = min(120, retry_delay * 1.5)
                                retry_delay += random.uniform(0, 0.5) * retry_delay
                        except:
                            retry_delay = min(120, retry_delay * 1.5)
                            
                        print(f"⏳ Rate limit hit, waiting {retry_delay:.2f} seconds before retry...")
                        time.sleep(retry_delay)
                        continue
                        
                    if "quota_exceeded" in error_msg.lower():
                        print(f"⚠️ Quota exceeded for segment {i+1}, waiting longer...")
                        time.sleep(180)
                        continue
                        
                    if attempt == max_retries - 1:
                        return (i, f"[Segment {i+1} error after {max_retries} attempts: {error_msg}]")
                    
                time.sleep(2)
                
            return (i, f"[Segment {i+1} failed after {max_retries} attempts]")
            
        finally:
            active_threads -= 1
            print(f"🏁 Finished segment {i+1} (Active threads: {active_threads})")

    # Create a single ThreadPoolExecutor for all segments
    max_workers = min(6, len(segments))  # Maximum of 6 concurrent workers
    print(f"\n🚀 Starting transcription with {max_workers} concurrent workers")
    
    with concurrent.futures.ThreadPoolExecutor(max_workers=max_workers) as executor:
        # Submit all segments for processing
        future_to_segment = {
            executor.submit(process_segment, (i, segment)): (i, segment)
            for i, segment in enumerate(segments)
        }
        
        # Process results as they complete
        results = []
        for future in concurrent.futures.as_completed(future_to_segment):
            try:
                result = future.result()
                results.append(result)
                
                if result[1].startswith("[Segment") and "error" in result[1]:
                    failed_segments.append(result[0] + 1)
                    
            except Exception as e:
                print(f"💥 Unexpected error processing segment: {str(e)}")
                segment_info = future_to_segment[future]
                results.append((segment_info[0], f"[Segment {segment_info[0]+1} unexpected error: {str(e)}]"))
                failed_segments.append(segment_info[0] + 1)
    
    # Sort results by segment index and create final transcript
    results.sort(key=lambda x: x[0])
    full_transcript = [text for _, text in results]
    
    print(f"\n📊 Parallel Processing Statistics:")
    print(f"Maximum concurrent threads: {max_active_threads}")
    
    if failed_segments:
        print("\n⚠️ Warning: Some segments failed to transcribe:")
        print(f"Failed segments: {sorted(failed_segments)}")
        print("Consider retrying these segments manually or with a different API key")
    else:
        print("\n✅ All segments transcribed successfully!")
    
    return full_transcript

def retry_with_backoff(func, max_retries=5, initial_delay=1):
    """Fonction utilitaire pour réessayer une opération avec un délai exponentiel"""
    def wrapper(*args, **kwargs):
        delay = initial_delay
        last_exception = None
        
        for attempt in range(max_retries):
            try:
                return func(*args, **kwargs)
            except Exception as e:
                last_exception = e
                error_code = str(e)
                if "429" in error_code or "499" in error_code: 
                    print(f"⚠️ Erreur API ({error_code}), nouvelle tentative {attempt + 1}/{max_retries} dans {delay} secondes...")
                    time.sleep(delay)
                    delay *= 2
                else:
                    raise e
        
        print(f"❌ Échec après {max_retries} tentatives : {str(last_exception)}")
        return None
    
    return wrapper

def create_word_pv_document(pv_text: str, meeting_info: dict) -> io.BytesIO:
    """Creates a Word document from PV text and meeting information."""
    doc = Document()

    # Get the first section
    section = doc.sections[0]

    # Add the standard header
    standard_header = """TANGER MED PORT AUTHORITY S.A "TMPA"
SOCIÉTÉ ANONYME À CONSEIL D'ADMINISTRATION
AU CAPITAL DE 1.704.000.000 DIRHAMS CONVERTIBLES
SIÈGE SOCIAL : ZONE FRANCHE DE KSAR EL MAJAZ, OUED RMEL,
COMMUNE ANJRA ROUTE DE FNIDEQ – TANGER
RC N°45349 TANGER – ICE : 000053443000022"""

    # Add a table for the header to apply border
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.text = standard_header

    # Center the text in the header cell
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Apply black border to the table
    from docx.oxml import ns as oxml_ns, OxmlElement

    tbl = table._tbl # get xml element of table
    tblPr = tbl.tblPr # get table properties element

    # Add border properties
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        border_elm = OxmlElement(f'w:tblBorders')
        border_elm.set(oxml_ns.qn('w:val'), 'single')
        border_elm.set(oxml_ns.qn('w:sz'), '15') # 1.5 points (value is in eighths of a point)
        border_elm.set(oxml_ns.qn('w:space'), '0')
        tblBorders.append(border_elm)
    tblPr.append(tblBorders)

    # Add some space after the header table
    doc.add_paragraph()

    # === Ajout du reste du texte généré par Gemini au corps du document ===
    # Parse the generated text to conditionally add sections based on content
    body_text = pv_text.strip()
    lines = body_text.split('\n')

    section_titles = [
        "PROCES VERBAL DE LA RÉUNION", # Special case: first major block
        "SONT PRESENTS OU REPRESENTES :",
        "Est Absent Excusé :",
        "Assistent également à la réunion :",
        "ORDRE DU JOUR:",
        "DÉROULÉ ET DÉCISIONS",
        "CONCLUSION",
        "ACRONYMES"
    ]

    # Function to check if a list of lines contains meaningful content
    def has_meaningful_content(lines_to_check, is_participant_list=False):
        if is_participant_list:
             # For participant lists, check specifically for bullet points with text
             return any(line.strip().startswith('-') and len(line.strip()) > 1 for line in lines_to_check)
        else:
             # For other sections, check for any non-empty line that isn't just a bullet point or asterisk
             return any(line.strip() and not line.strip().startswith('-') and not line.strip().startswith('*') for line in lines_to_check)

    # Function to add a section (title + content) to the document if it has content
    def add_section_to_doc(title, content_lines):
        is_participant_list = title in ["SONT PRESENTS OU REPRESENTES :", "Est Absent Excusé :", "Assistent également à la réunion :"]

        # Always include the main PV header section if it's the first major block
        if title == "PROCES VERBAL DE LA RÉUNION  DU CONSEIL D'ADMINISTRATION ":
             # Add the lines of the main PV header block with specific formatting for the first three lines
             for i, line in enumerate(content_lines):
                 stripped_line = line.strip()
                 if stripped_line:
                     paragraph = doc.add_paragraph()
                     run = paragraph.add_run(stripped_line)
                     if i < 3: # Apply bold and center to the first three lines
                         run.bold = True
                         paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                     # No specific formatting for lines beyond the second one in this block

             doc.add_paragraph() # Add space after this block
             return
             
        # For other sections, check if there is meaningful content before adding the title and content
        if has_meaningful_content(content_lines, is_participant_list):
            # Add title as a bold paragraph
            title_para = doc.add_paragraph()
            title_run = title_para.add_run(title)
            title_run.bold = True # Ensure this is always applied when the condition above is met

            # Re-apply bold to be sure (should not be necessary but for robustness)
            title_run.font.bold = True

            # List of problematic phrases to exclude if they appear as a full line
            problematic_phrases = [
                "Le PV a été validé moyennant les corrections à apporter.",
                "Aucune conclusion formelle enregistrée.", # Fallback phrase for Conclusion
                "Aucun acronyme n'a été trouvé dans les documents fournis.", # Fallback phrase for Acronymes
                # Add other phrases here if they appear in future generations
            ]

            # Add the content lines as paragraphs, skipping problematic phrases
            for line in content_lines:
                stripped_line = line.strip()
                if stripped_line:
                    # Check if the stripped line is one of the problematic phrases
                    is_problematic = any(stripped_line == phrase for phrase in problematic_phrases)
                    
                    if not is_problematic:
                        doc.add_paragraph(stripped_line)
                        
            doc.add_paragraph() # Add a space after the section


    current_section_title = None
    current_section_lines = []

    # Process lines from the generated text
    for i, line in enumerate(lines):
        stripped_line = line.strip()
        
        # Check if the current line is a known section title
        is_title = False
        for title in section_titles:
            if stripped_line.startswith(title):
                is_title = True
                # If we were processing a previous section, add it to the document now
                if current_section_title is not None:
                    # Pass content lines (excluding the title line if it was included in collected lines)
                    add_section_to_doc(current_section_title, current_section_lines)
                
                # Start a new section: set the current title and initialize collected lines
                current_section_title = title
                current_section_lines = [] # Start with an empty list for the new section's content
                # Note: The title line itself is not added to current_section_lines as content
                break

        # If the line is not a section title, add it to the current section's collected lines
        if not is_title:
            # Only add lines if we are currently inside a section (after the first title)
            if current_section_title is not None:
                 current_section_lines.append(line)
            else: # Handle any text before the first recognized section title
                 # This part might include the main PV header block if it wasn't stripped earlier
                 # Add these lines directly if they have content, before any titled sections begin.
                 if stripped_line:
                      doc.add_paragraph(stripped_line)

    # After the loop, add the last section if there was one being processed
    if current_section_title is not None:
        add_section_to_doc(current_section_title, current_section_lines)

    # === Add Fixed Closing Text and Signatures ===
    doc.add_paragraph("Le Conseil d'Administration confère tous pouvoirs au porteur d'un original, d'une copie ou d'un extrait du présent procès-verbal aux fins d’accomplir toutes les formalités requises par la loi. ") # Add first closing paragraph
    doc.add_paragraph("Plus rien n'étant à l'ordre du jour et personne ne demandant la parole, le Président remercie l’ensemble des membres du Conseil d’Administration et déclare que la séance est levée. ") # Add second closing paragraph
    doc.add_paragraph("De tout ce que dessus, il a été établi le présent procès-verbal pour servir et valoir ce que de droit.") # Add third closing paragraph
    doc.add_paragraph("Etabli en trois (3) exemplaires originaux") # Add fourth closing paragraph

    # Add space before signatures
    doc.add_paragraph()
    doc.add_paragraph()

    # Add signature lines (centered)
    president_para = doc.add_paragraph("Président du Conseil d'Administration")
    president_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add space for signature
    doc.add_paragraph()
    doc.add_paragraph()

    administrator_para = doc.add_paragraph("Administrateur")
    administrator_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add space for signature
    doc.add_paragraph()
    doc.add_paragraph()

    secretary_para = doc.add_paragraph("Secrétaire du Conseil d'Administration")
    secretary_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # === Numérotation des pages ===
    # Get the footer
    footer = section.footer

    # Add standard footer text
    date_for_footer = meeting_info.get('date', '').replace('/', '_').replace('-', '_')
    footer_text = f"PV_CA_TMPA_{date_for_footer}"
    footer_para = footer.add_paragraph()
    footer_para.text = footer_text

    # Add a tab to separate text and page number
    footer_para.add_run('\t')

    # Add page number field
    run = footer_para.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(oxml_ns.qn('w:fldCharType'), 'begin')
    run._r.append(fldChar)
    instrText = OxmlElement('w:instrText')
    instrText.text = 'PAGE \\\* MERGEFORMAT' # Field code for page number
    run._r.append(instrText)
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(oxml_ns.qn('w:fldCharType'), 'end')
    run._r.append(fldChar)

    # Add total pages field (optional, but common)
    run = footer_para.add_run(" sur ")
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(oxml_ns.qn('w:fldCharType'), 'begin')
    run._r.append(fldChar)
    instrText = OxmlElement('w:instrText')
    instrText.text = 'NUMPAGES \\\* MERGEFORMAT' # Field code for total pages
    run._r.append(instrText)
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(oxml_ns.qn('w:fldCharType'), 'end')
    run._r.append(fldChar)

    # Set up tab stops for left and right alignment on the same line
    from docx.shared import Inches
    from docx.text.tabstops import TabStop
    from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER

    # Clear existing tab stops
    footer_para.paragraph_format.tab_stops.clear_all()

    # Add a right-aligned tab stop at the right margin (adjust position as needed)
    # Assuming default page width and margins, approx 6.5 inches from left margin
    right_tab_pos = Inches(6.5) # You might need to adjust this based on your document layout
    footer_para.paragraph_format.tab_stops.add_tab_stop(
        right_tab_pos,
        WD_TAB_ALIGNMENT.RIGHT,
        WD_TAB_LEADER.SPACES
    )

    # The default alignment of a paragraph is left, so the footer text will be left-aligned
    # The page number will align to the right_tab_pos after the tab character
    footer_para.alignment = WD_ALIGN_PARAGRAPH.LEFT # Ensure paragraph is left-aligned

    # Save the document to a BytesIO object
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def process_handwritten_image(image_bytes):
    """Extrait le texte d'une image manuscrite avec mécanisme de retry"""
    @retry_with_backoff
    def transcribe_image():
        try:
            image_base64 = base64.b64encode(image_bytes).decode('utf-8')
            
            model = genai.GenerativeModel('gemini-2.0-flash')
            
            prompt = """Transcris précisément le texte manuscrit dans cette image.
            INSTRUCTIONS :
            1. Retourne uniquement le texte, sans commentaires
            2. Préserve la mise en forme (retours à la ligne, espacements)
            3. Conserve la ponctuation exacte
            4. Maintiens les nombres et symboles tels quels
            5. Respecte les majuscules et minuscules"""
            
            response = model.generate_content([
                prompt,
                {"mime_type": "image/jpeg", "data": image_base64}
            ])
            
            if response.text:
                return response.text.strip()
            else:
                raise Exception("Aucun texte détecté dans l'image.")
                
        except Exception as e:
            print(f"⚠️ Tentative de transcription échouée : {str(e)}")
            raise e

    try:
        # Premier essai
        result = transcribe_image()
        if result:
            return result
            
        # Si le résultat est vide, on attend et on réessaie
        time.sleep(2)
        print("🔄 Nouvelle tentative de transcription...")
        
        # Deuxième essai avec un prompt plus détaillé
        prompt_retry = """Analyse et transcris TOUT le texte manuscrit visible dans cette image.
        IMPORTANT :
        - Examine l'image en détail, pixel par pixel
        - Transcris absolument tout le texte visible
        - N'oublie aucun détail, même les petites annotations
        - Conserve la structure exacte du texte
        - Inclus les numéros, symboles et caractères spéciaux"""
        
        model = genai.GenerativeModel('gemini-2.0-flash')
        image_base64 = base64.b64encode(image_bytes).decode('utf-8')
        
        response = model.generate_content([
            prompt_retry,
            {"mime_type": "image/jpeg", "data": image_base64}
        ])
        
        if response.text:
            return response.text.strip()
        else:
            print("⚠️ Aucun texte n'a été détecté dans l'image après plusieurs tentatives.")
            return ""
            
    except Exception as e:
        print(f"❌ Erreur lors de la reconnaissance du texte : {str(e)}")
        return ""

def process_pdf(pdf_bytes):
    """Extrait le contenu détaillé et les acronymes d'un PDF en un seul appel."""
    try:
        pdf_base64 = base64.b64encode(pdf_bytes).decode('utf-8')
        
        model = genai.GenerativeModel('gemini-2.0-flash')
        
        prompt = """Analyse ce document PDF de manière EXHAUSTIVE et DÉTAILLÉE.
        
        INSTRUCTIONS SPÉCIFIQUES :
        
        1. EXTRACTION COMPLÈTE DU CONTENU :
           - Extraire TOUS les textes, exactement comme ils apparaissent.
           - Conserver TOUS les chiffres, statistiques, données numériques avec leurs unités.
           - Maintenir TOUS les tableaux avec leurs données complètes.
           - Décrire TOUS les graphiques avec leurs valeurs précises.
           - Capturer TOUTES les notes de bas de page et références.
           - Respecter la structure (sections, titres, listes).
           - NE PAS résumer ou synthétiser le corps du texte.
           
        2. EXTRACTION DES ACRONYMES :
           - Identifier TOUS les acronymes présents dans le document.
           - Si l'acronyme est défini explicitement dans le texte, utiliser cette définition EXACTE.
           - Si l'acronyme n'est PAS défini dans le texte, rechercher sa définition officielle connue dans des sources fiables.
           - Lister les acronymes et leurs définitions SÉPARÉMENT à la fin.
        
        3. FORMAT DE SORTIE ATTENDU :
           - D'abord, le contenu complet et détaillé du document, en respectant sa structure.
           - Ensuite, une ligne de séparation claire comme : '--- ACRONYMES ---'.
           - Enfin, la liste des acronymes, un par ligne, au format : 'ACRONYME: Définition complète'.
           
        IMPORTANT : Assure-toi de bien séparer le contenu principal de la liste des acronymes avec '--- ACRONYMES ---'."""
        
        @retry_with_backoff
        def analyze_pdf_and_extract_acronyms():
            response = model.generate_content([
                {
                    "role": "user",
                    "parts": [
                        prompt,
                        {"mime_type": "application/pdf", "data": pdf_base64}
                    ]
                }
            ])
            return response.text if response.text else ""
        
        full_result = analyze_pdf_and_extract_acronyms()
        
        if not full_result:
            print(f"⚠️ Aucun contenu extrait du PDF")
            return {"summary": "", "acronyms": {}}
            
        # Séparer le contenu et les acronymes
        separator = "--- ACRONYMES ---"
        if separator in full_result:
            summary_part, acronym_part = full_result.split(separator, 1)
            summary = summary_part.strip()
            
            # Parser les acronymes
            acronyms = {}
            lines = acronym_part.strip().split('\n')
            for line in lines:
                if ':' in line:
                    acronym, definition = line.split(':', 1)
                    acronym = acronym.strip().upper()
                    definition = definition.strip()
                    if acronym and definition:
                        acronyms[acronym] = definition
            return {"summary": summary, "acronyms": acronyms}
        else:
            # Si le séparateur n'est pas trouvé, retourner tout comme résumé et pas d'acronymes
            print(f"⚠️ Séparateur d'acronymes non trouvé dans l'analyse")
            return {"summary": full_result.strip(), "acronyms": {}}
            
    except Exception as e:
        print(f"❌ Erreur lors de l'analyse du PDF: {str(e)}")
        return {"summary": f"[Erreur lors de l'analyse du PDF: {str(e)}]", "acronyms": {}}

# --- Dependencies ---

async def require_video_or_audio(
    video: Optional[UploadFile] = File(None),
    audio: List[UploadFile] = File([]),
    meetingData: str = Form(...),
):
    """Dependency to ensure at least one video or audio file or Google Drive URL is provided."""
    # Parse meetingData to access googleDriveUrl
    try:
        meeting_info = json.loads(meetingData)
        google_drive_url = meeting_info.get("googleDriveUrl")
    except json.JSONDecodeError:
        # If meetingData is invalid, the main endpoint will handle it, 
        # but for this dependency, we can assume it's not the required media source.
        google_drive_url = None

    if not video and not audio and not google_drive_url:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="At least one video or audio file or Google Drive URL is required for PV generation."
        )

# --- PV Text Generation Function ---
async def generate_pv_text_with_gemini(
    meeting_info: dict,
    video_transcript: str,
    audio_transcripts_list: List[str],
    ocr_texts_list: List[str],
    pdf_results_list: List[dict]
) -> str:
    """Generates structured PV text using Gemini based on processed media content and meeting info."""
    try:
        # Combine all processed text sources into a single string for the prompt
        combined_text = ""

        if video_transcript:
            combined_text += "[TRANSCRIPTION VIDÉO]\n" + video_transcript.strip() + "\n\n"

        if audio_transcripts_list:
            combined_text += "[ENREGISTREMENTS AUDIO]\n"
            combined_text += "\n---\n".join([t.strip() for t in audio_transcripts_list if t.strip()]) + "\n\n"

        if ocr_texts_list:
            combined_text += "[NOTES MANUSCRITES (OCR)]\n"
            combined_text += "\n---\n".join([t.strip() for t in ocr_texts_list if t.strip()]) + "\n\n"

        if pdf_results_list:
            combined_text += "[DOCUMENTS PDF]\n"
            # Combine summaries
            combined_pdf_summaries = "\n---\n".join([res["summary"].strip() for res in pdf_results_list if res and "summary" in res and res["summary"].strip()])
            if combined_pdf_summaries:
                 combined_text += "## Résumés :\n" + combined_pdf_summaries + "\n\n"

            # Combine acronymsimage.png
            all_acronyms = {}
            for res in pdf_results_list:
                 if res and "acronyms" in res:
                     all_acronyms.update(res["acronyms"])

            if all_acronyms:
                 combined_text += "## Acronymes :\n"
                 for acronym, definition in all_acronyms.items():
                      combined_text += f"{acronym}: {definition}\n"
            combined_text += "\n"

        if not combined_text.strip():
            return "Aucun contenu médiatique traité pour générer le PV."

        # Construct the prompt for Gemini, based on the desired PV structure
        # We will build the participant lists first
        participants = meeting_info.get('participants', [])

        present_participants = [p['nom'] for p in participants if p.get('statut') == 'Present']
        absent_participants = [p['nom'] for p in participants if p.get('statut') == 'Absent Excusé']
        assistant_participants = [p['nom'] for p in participants if p.get('statut') == 'Assistant']

        # Format the lists for the prompt (using bullet points)
        present_participants_list = "\n".join([f"- {name}" for name in present_participants]) if present_participants else ""
        absent_participants_list = "\n".join([f"- {name}" for name in absent_participants]) if absent_participants else ""
        assistant_participants_list = "\n".join([f"- {name}" for name in assistant_participants]) if assistant_participants else ""

        # Construct the prompt for Gemini, combining new participant logic with old prompt style
        full_prompt_content = f"""PROCES VERBAL DE LA RÉUNION  DU CONSEIL D'ADMINISTRATION  
DU {meeting_info.get('date', 'N/A')}
À {meeting_info.get('time', 'N/A')} heures.

En  {meeting_info.get('year', 'N/A')} à écrire en toutes lettres pour donner comme cet example : L'An Deux Mille Vingt-Cinq, ,  
le {meeting_info.get('date', 'N/A')} à écrire en toutes lettres pour donner comme cet example : Le Dix Février,  
à {meeting_info.get('time', 'N/A')} à écrire en toutes lettres pour donner comme cet example : À 15 heures.  

Les membres du Conseil d'Administration de Tanger Med Port Authority S.A, par abréviation,
" TMPA " se sont réunis en Conseil d'Administration en présentiel au bureau de {meeting_info.get('location', '')} et par visioconférence conformément aux dispositions réglementaires.


SONT PRESENTS OU REPRESENTES :
{present_participants_list}

Est Absent Excusé :
{absent_participants_list}

Assistent également à la réunion :
{assistant_participants_list}


ORDRE DU JOUR:
[Lister ici les points de l'ordre du jour numérotés, extraits du contenu traité. Utiliser une liste numérotée comme dans l'exemple.]


DÉROULÉ ET DÉCISIONS

[Pour chaque point de l'ordre du jour listé ci-dessus, fournir un résumé détaillé basé sur le 'Contenu Traité Brut'. Inclure les discussions, les décisions prises et les résolutions. Structurez cela clairement point par point.]

[POINT N°] [Titre du point]
[Résumé des discussions et points clés abordés, basé sur le Contenu Traité Brut]
Décisions : [Décisions spécifiques prises pour ce point, basées sur le Contenu Traité Brut]
Résolutions : [Résolutions spécifiques adoptées pour ce point, basées sur le Contenu Traité Brut]

[Répéter pour chaque point de l'ordre du jour]

CONCLUSION
[Résumer ici les principaux aboutissements de la réunion, les décisions importantes prises, et les éventuelles prochaines étapes ou actions à entreprendre, basé sur le Contenu Traité Brut.]

ACRONYMES
[Lister ici les acronymes identifiés et leurs définitions complètes, extraits spécifiquement de la partie Acronymes des résultats PDF, si disponibles.]


Contenu Traité Brut (pour référence interne uniquement, ne pas inclure ceci dans le PV final):
{combined_text}

INSTRUCTIONS POUR LA GÉNÉRATION DU PV :
1. Le texte généré DOIT suivre la structure définie ci-dessus, incluant les sections "PROCES VERBAL...", "ORDRE DU JOUR", "Présents", "Absents excusés", "Assistent également", " DÉROULÉ ET DÉCISIONS ", " CONCLUSION ", et " ACRONYMES".
2. Remplir les sections du PV EN UTILISANT STRICTEMENT UNIQUEMENT les informations pertinentes extraites du 'Contenu Traité Brut'.
3. Pour les sections "Présents", "Absents excusés", "Assistent également", utiliser les listes de participants fournies directement dans le prompt , si l;une est vide supprime la du Pv genere.
4. Pour l'ORDRE DU JOUR, lister les points tels qu'ils apparaissent ou sont déduits du 'Contenu Traité Brut'. Utiliser une liste numérotée (ex: 1., 2., ...).
5. Pour le DÉROULÉ ET DÉCISIONS, parcourir l'ordre du jour et résumer les discussions, décisions, et résolutions pour chaque point, en se basant EXCLUSIVEMENT sur le 'Contenu Traité Brut'. Commencer chaque point par le numéro et le titre (ex: [POINT N°] [Titre du point]), suivi des sous-sections (Discussions, Décisions, Résolutions) si l'information est présente dans le contenu et si jamais y'a autre chose d'important a citer c'est a citer .
6. Pour la CONCLUSION, extraire les éléments de conclusion et les prochaines étapes du 'Contenu Traité Brut'.
7. Pour les ACRONYMES, lister UNIQUEMENT ceux qui ont été extraits et fournis dans la section [DOCUMENTS PDF] du 'Contenu Traité Brut'. Si aucun acronyme n'est fourni dans cette section, pas besoin de la citer.
8. Maintenir un ton professionnel et formel, caractéristique d'un procès-verbal officiel.
9. NE PAS inclure la section "Contenu Traité Brut" ou les "INSTRUCTIONS POUR LA GÉNÉRATION DU PV" dans le texte final du PV. Elles sont fournies uniquement pour  générer le texte correct.

⚠️ Ne pas écrire de "N/A". Si une information est manquante, ignorer ou laisser vide.
⚠️ Ne pas conserver les crochets [], remplacer par titres clairs.
⚠️ Organise le PV de manière propre et professionnelle avec des titres hiérarchisés.
10. ⚠️ Si une section ne contient aucune information pertinente extraite du 'Contenu Traité Brut' (ex. : ACRONYMES, QUESTIONS DIVERSES, CONCLUSION), alors :
- Ne pas la générer.
- Ne pas insérer de titre vide.
- Ne pas écrire de phrase du type "Aucune information disponible".
- Supprimer la section entière du PV final.
11. ⚠️ Si un placeholder comme  [Titre du point], etc. ne peut pas être remplacé par une donnée réelle du contenu traité, alors :
- Supprimer toute la phrase contenant ce placeholder.
-  Ne pas afficher le placeholder dans le texte final.
"""

        model = genai.GenerativeModel('gemini-2.0-flash')

        @retry_with_backoff
        async def call_gemini_for_pv():
            print("Attempting Gemini call for PV generation...") # Debug print
            response = model.generate_content(
                [{"role": "user", "parts": [full_prompt_content]}],  # Pass prompt as parts in a user role
                request_options={"timeout": 180} # Increased timeout
            )
            print(f"Gemini PV generation response status: {response.candidates[0].finish_reason if response.candidates else 'No candidates'}") # Debug print
            return response.text if response.text else ""

        generated_text = await call_gemini_for_pv()

        if not generated_text or not generated_text.strip():
            print("⚠️ Gemini generated empty PV text.") # Debug print
            return "[Échec de la génération de texte de PV par IA ou texte vide]"

        # Post-processing: Remove Raw Processed Content and Instructions sections by splitting
        raw_content_tag = "Contenu Traité Brut (pour référence interne uniquement, ne pas inclure ceci dans le PV final):"
        instructions_tag = "INSTRUCTIONS POUR LA GÉNÉRATION DU PV :"

        if raw_content_tag in generated_text:
            generated_text = generated_text.split(raw_content_tag, 1)[0].strip()

        if instructions_tag in generated_text:
             generated_text = generated_text.split(instructions_tag, 1)[0].strip()

        # Basic post-processing (remove common markdown formatting and any remaining instructions)
        generated_text = generated_text.replace('**', '')
        generated_text = generated_text.replace('*', '')

        # Clean up multiple newlines and leading/trailing whitespace
        generated_text = re.sub(r'\n{2,}', '\n\n', generated_text).strip()

        # Remove any remaining specific instruction placeholders that weren't handled by the split
        # Be cautious with broad regex; target specific patterns if possible.
        instruction_placeholders_regex = [
             r'\[Lister ici les points de l\'ordre du jour numérotés, extraits du contenu traité\. Utiliser une liste numérotée comme dans l\'exemple\.\]',
             r'\[Pour chaque point de l\'ordre du jour listé ci-dessus, fournir un résumé détaillé basé sur le \'Contenu Traité Brut\'\. Inclure les discussions, les décisions prises et les résolutions\. Structurez cela clairement point par point\.\]',
             r'\[POINT N°\] \[Titre du point\]\n\[Résumé des discussions et points clés abordés, basé sur le Contenu Traité Brut\]\nDécisions : \[Décisions spécifiques prises pour ce point, basées sur le Contenu Traité Brut\]\nRésolutions : \[Résolutions spécifiques adoptées pour ce point, basées sur le Contenu Traité Brut\]\n\[Répéter pour chaque point de l\'ordre du jour\]',
             r'\[Résumer ici les principaux aboutissements de la réunion, les décisions importantes prises, et les éventuelles prochaines étapes ou actions à entreprendre, basé sur le Contenu Traité Brut\.\]',
             r'\[Lister ici les acronymes identifiés et leurs définitions complètes, extraits spécifiquement de la partie Acronymes des résultats PDF, si disponibles\.\]',
             r'\[Lister ici les participants présents, extraits du contenu traité ou des informations de la réunion\.\]',
             r'\[Lister ici les absents excusés, extraits du contenu traité\.\]',
             r'\[Lister ici les personnes assistant à la réunion \(invités, secrétariat, etc\.\), extraits du contenu traité\.\]',
        ]

        for regex_pattern in instruction_placeholders_regex:
             generated_text = re.sub(regex_pattern, '', generated_text, flags=re.DOTALL | re.IGNORECASE).strip()

        # Clean up again after removing placeholders
        generated_text = re.sub(r'\n{2,}', '\n\n', generated_text).strip()

        print("PV text generation completed by Gemini.") # Debug print
        return generated_text.strip()

    except Exception as e:
        print(f"❌ Error during PV text generation: {str(e)}") # Debug print
        return f"[Erreur lors de la génération du texte du PV : {str(e)}]"

# --- API Endpoints ---

@app.post("/transcribe_video")
async def transcribe_video(
    video: Optional[UploadFile] = File(None),
    drive_url: Optional[str] = Form(None)
):
    """Full video transcription pipeline. Accepts file upload or Google Drive link."""
    print(f"Received request - video: {video}, drive_url: {drive_url}")
    
    with tempfile.TemporaryDirectory() as temp_dir:
        try:
            # 1. Handle file upload or Google Drive link
            if video is not None:
                print(f"Processing uploaded video: {video.filename}, size: {video.size} bytes")
                ext = os.path.splitext(video.filename)[1].lower() if video.filename else '.mp4'
                video_temp_path = os.path.join(temp_dir, f"uploaded_video{ext}")
                written_size = 0
                with open(video_temp_path, 'wb') as out_file:
                    while True:
                        chunk = await video.read(1024 * 1024)
                        if not chunk:
                            break
                        out_file.write(chunk)
                        written_size += len(chunk)
                print(f"Video saved to: {video_temp_path}, written size: {written_size} bytes")
            elif drive_url:
                print(f"Processing drive URL: {drive_url}")
                video_temp_path = os.path.join(temp_dir, "downloaded_video.mp4")
                ok, err = download_video_from_drive(drive_url, video_temp_path)
                if not ok:
                    print(f"Drive download failed: {err}")
                    return JSONResponse(status_code=400, content={"error": err})
            else:
                print("No video file or drive_url provided")
                return JSONResponse(status_code=400, content={"error": "No video file or drive_url provided."})
            
            # 2. Verify video
            print("Verifying video file...")
            valid, err = verify_video_file(video_temp_path)
            if not valid:
                print(f"Video verification failed: {err}")
                return JSONResponse(status_code=400, content={"error": err})
            
            # 3. Extract audio
            print("Extracting audio...")
            audio_path = os.path.join(temp_dir, "output_audio.mp3")
            ok, err = extract_audio_from_video(video_temp_path, audio_path)
            if not ok:
                print(f"Audio extraction failed: {err}")
                return JSONResponse(status_code=400, content={"error": err})
            
            # 4. Segment audio
            print("Segmenting audio...")
            segments = segment_audio(audio_path)
            if not segments:
                print("Audio segmentation failed")
                return JSONResponse(status_code=400, content={"error": "Audio segmentation failed."})
            
            # 5. Transcribe segments
            print("Transcribing segments...")
            print("Video temp path:", video_temp_path)
            print("Audio path:", audio_path)
            print("Number of segments:", len(segments))
            transcript_segments = transcribe_audio_segments(segments)
            transcript = "\n".join(transcript_segments)
            print("Transcription completed successfully")
            return {"transcript": transcript}
            
        except Exception as e:
            print(f"Error in transcribe_video: {str(e)}")
            return JSONResponse(status_code=500, content={"error": str(e)})

@app.post("/transcribe_audio")
async def transcribe_audio(audio: Optional[UploadFile] = File(None)):
    """Pipeline complet de transcription pour un fichier audio (uploadé ou enregistré)"""
    with tempfile.TemporaryDirectory() as temp_dir:
        try:
            if audio is None:
                return JSONResponse(status_code=400, content={"error": "Aucun fichier audio fourni."})
            
            ext = os.path.splitext(audio.filename)[1].lower() if audio.filename else '.mp3'
            audio_temp_path = os.path.join(temp_dir, f"uploaded_audio{ext}")

            # 1. Sauvegarder le fichier audio temporairement
            with open(audio_temp_path, 'wb') as out_file:
                while True:
                    chunk = await audio.read(1024 * 1024)
                    if not chunk:
                        break
                    out_file.write(chunk)

            # 2. Convertir en MP3 si besoin
            if ext != ".mp3":
                converted_audio_path = os.path.join(temp_dir, "converted_audio.mp3")
                convert_cmd = [
                    "ffmpeg", "-y", "-i", audio_temp_path,
                    "-acodec", "libmp3lame", "-ar", "44100", "-b:a", "192k",
                    converted_audio_path
                ]
                subprocess.run(convert_cmd, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
                audio_path = converted_audio_path
            else:
                audio_path = audio_temp_path

            # 3. Segmenter
            segments = segment_audio(audio_path)
            if not segments:
                return JSONResponse(status_code=400, content={"error": "Échec de la segmentation audio."})

            # 4. Transcrire
            print("Audio path:", audio_path)
            print("Segments:", segments)
            
            transcript_segments = transcribe_audio_segments(segments)
            transcript = "\n".join(transcript_segments)

            return {"transcription": transcript}

        except Exception as e:
            return JSONResponse(status_code=500, content={"error": str(e)})


@app.post("/ocr_handwritten")
async def ocr_handwritten(images: List[UploadFile] = File(...)):
    """Transcrit le texte manuscrit à partir d'une ou bien plusieurs images."""
    try:
        results = {}
        for image in images:
            try:
                # Lire le contenu de l'image
                image_bytes = await image.read()
                
                # Traiter l'image
                transcription = process_handwritten_image(image_bytes)
                
                # Stocker le résultat
                results[image.filename] = {
                    "success": True,
                    "text": transcription
                }
                
            except Exception as e:
                results[image.filename] = {
                    "success": False,
                    "error": str(e)
                }
        
        return {"results": results}
        
    except Exception as e:
        return JSONResponse(
            status_code=500,
            content={"error": f"Erreur lors du traitement des images : {str(e)}"}
        )

@app.post("/extract_pdf")
async def extract_pdf(pdf: UploadFile = File(...)):
    """Extrait le contenu et les acronymes d'un PDF."""
    try:
        # Lire le contenu du PDF
        pdf_bytes = await pdf.read()
        
        # Traiter le PDF
        result = process_pdf(pdf_bytes)
        
        # Ajouter le nom du fichier au résultat
        result["filename"] = pdf.filename
        
        return result
        
    except Exception as e:
        return JSONResponse(
            status_code=500,
            content={"error": f"Erreur lors du traitement du PDF : {str(e)}"}
        )

@app.post("/generate_pv", dependencies=[Depends(require_video_or_audio)])
async def generate_pv(
    meetingData: str = Form(...),
    video: Optional[UploadFile] = File(None),
    audio: List[UploadFile] = File([]),
    images: List[UploadFile] = File([]),
    pdfs: List[UploadFile] = File([]),
):
    # Placeholder: implement PV generation logic

    # 1. Receive and parse meeting data
    try:
        meeting_info = json.loads(meetingData)
        print(f"Received meeting data: {meeting_info}") # Debug print
    except json.JSONDecodeError:
        return JSONResponse(status_code=400, content={"error": "Invalid meeting data format."})

    # Access Google Drive URL if present
    google_drive_url = meeting_info.get("googleDriveUrl")

    # Create a temporary directory to store uploaded files
    with tempfile.TemporaryDirectory() as temp_dir:
        print(f"Created temporary directory: {temp_dir}") # Debug print
        video_path = None
        audio_paths = []
        image_paths = []
        pdf_paths = []

        # 2. Save uploaded files to the temporary directory
        try:
            # Save Video File
            if video:
                video_filename = video.filename if video.filename else "video.mp4"
                video_path = os.path.join(temp_dir, video_filename)
                with open(video_path, 'wb') as f:
                    while True:
                        chunk = await video.read(1024 * 1024)
                        if not chunk:
                            break
                        f.write(chunk)
                print(f"Saved video to: {video_path}") # Debug print

            # Save Audio Files
            for i, audio_file in enumerate(audio):
                audio_filename = audio_file.filename if audio_file.filename else f"audio_{i}.mp3"
                audio_path = os.path.join(temp_dir, audio_filename)
                with open(audio_path, 'wb') as f:
                    while True:
                        chunk = await audio_file.read(1024 * 1024)
                        if not chunk:
                            break
                        f.write(chunk)
                audio_paths.append(audio_path)
                print(f"Saved audio file {i} to: {audio_path}") # Debug print

            # Save Image Files
            for i, image_file in enumerate(images):
                image_filename = image_file.filename if image_file.filename else f"image_{i}.png"
                image_path = os.path.join(temp_dir, image_filename)
                with open(image_path, 'wb') as f:
                    while True:
                        chunk = await image_file.read(1024 * 1024)
                        if not chunk:
                            break
                        f.write(chunk)
                image_paths.append(image_path)
                print(f"Saved image file {i} to: {image_path}") # Debug print

            # Save PDF Files
            for i, pdf_file in enumerate(pdfs):
                pdf_filename = pdf_file.filename if pdf_file.filename else f"pdf_{i}.pdf"
                pdf_path = os.path.join(temp_dir, pdf_filename)
                with open(pdf_path, 'wb') as f:
                    while True:
                        chunk = await pdf_file.read(1024 * 1024)
                        if not chunk:
                            break
                        f.write(chunk)
                pdf_paths.append(pdf_path)
                print(f"Saved PDF file {i} to: {pdf_path}") # Debug print

        except Exception as e:
            print(f"Error saving uploaded files: {str(e)}") # Debug print
            return JSONResponse(status_code=500, content={"error": f"Failed to save uploaded files: {str(e)}"})

        # --- Processing logic starts here ---
        print("Starting media processing...") # Debug print
        video_transcript = ""
        audio_transcripts_list = [] # Use a list to store transcripts from multiple audio files
        ocr_texts_list = [] # Use a list to store texts from multiple image files
        pdf_results_list = [] # Use a list to store results from multiple PDF files

        # Process Video (if uploaded) or Google Drive URL
        if video_path or google_drive_url:
            print("Processing video/Google Drive URL...") # Debug print
            try:
                 # Reuse logic from old transcribe_video endpoint
                 if google_drive_url:
                     # Download from Drive
                     downloaded_video_path = os.path.join(temp_dir, "downloaded_video.mp4")
                     ok, err = download_video_from_drive(google_drive_url, downloaded_video_path)
                     if not ok:
                         print(f"Drive download failed: {err}") # Debug print
                         # Decide if this is a critical error or if we can continue with other media
                         # For now, let's add an error placeholder to the transcript
                         video_transcript = f"[Erreur de téléchargement Google Drive: {err}]"
                         video_to_process_path = None # No video file to process further
                     else:
                          video_to_process_path = downloaded_video_path
                          print(f"Downloaded video to: {video_to_process_path}") # Debug print
                 else:
                     # Use the uploaded video file
                     video_to_process_path = video_path

                 if video_to_process_path:
                     # Verify video
                     valid, err = verify_video_file(video_to_process_path)
                     if not valid:
                         print(f"Video verification failed: {err}") # Debug print
                         video_transcript = f"[Erreur de vérification vidéo: {err}]"
                     else:
                          # Extract audio
                          audio_from_video_path = os.path.join(temp_dir, "audio_from_video.mp3")
                          ok, err = extract_audio_from_video(video_to_process_path, audio_from_video_path)
                          if not ok:
                              print(f"Audio extraction failed: {err}") # Debug print
                              video_transcript = f"[Erreur d'extraction audio vidéo: {err}]"
                          else:
                               # Segment and transcribe audio from video
                               segments = segment_audio(audio_from_video_path)
                               if not segments:
                                   print("Audio segmentation failed for video") # Debug print
                                   video_transcript = "[Échec de la segmentation audio vidéo]"
                               else:
                                   transcript_segments = transcribe_audio_segments(segments)
                                   video_transcript = "\n".join(transcript_segments)
                                   print("Video transcription completed.") # Debug print

            except Exception as e:
                 print(f"Error processing video/Google Drive URL: {str(e)}") # Debug print
                 video_transcript = f"[Erreur de traitement vidéo/URL: {str(e)}]"

        # Process Audio Files
        if audio_paths:
            print(f"Processing {len(audio_paths)} audio file(s)...") # Debug print
            for i, audio_file_path in enumerate(audio_paths):
                try:
                    # Reuse logic from old transcribe_audio endpoint (saving, converting, segmenting, transcribing)
                    # Note: saving is already done, so start from convert if needed
                    
                    # Check extension and convert if needed (simplified here, could be more robust)
                    ext = os.path.splitext(audio_file_path)[1].lower()
                    processed_audio_path = audio_file_path
                    if ext not in [".mp3", ".wav", ".aac", ".flac", ".m4a"]:
                        # Simple conversion attempt - might need more specific logic
                        converted_audio_path = os.path.join(temp_dir, f"converted_audio_{i}.mp3")
                        convert_cmd = [
                            "ffmpeg", "-y", "-i", audio_file_path,
                            "-acodec", "libmp3lame", "-ar", "44100", "-b:a", "192k",
                            converted_audio_path
                        ]
                        # Note: In a real app, you might want to check return code and handle errors
                        subprocess.run(convert_cmd, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
                        processed_audio_path = converted_audio_path
                        print(f"Converted audio file {i} to MP3: {processed_audio_path}") # Debug print
                    
                    segments = segment_audio(processed_audio_path)
                    if not segments:
                        print(f"Audio segmentation failed for file {i}") # Debug print
                        audio_transcripts_list.append(f"[Échec de la segmentation audio fichier {i}]")
                    else:
                        transcript_segments = transcribe_audio_segments(segments)
                        audio_transcripts_list.append("\n".join(transcript_segments))
                        print(f"Audio transcription completed for file {i}.") # Debug print

                except Exception as e:
                     print(f"Error processing audio file {i}: {str(e)}") # Debug print
                     audio_transcripts_list.append(f"[Erreur de traitement audio fichier {i}: {str(e)}]")

        # Process Image Files
        if image_paths:
            print(f"Processing {len(image_paths)} image file(s)...") # Debug print
            for i, image_file_path in enumerate(image_paths):
                try:
                    # Read image bytes from saved file
                    with open(image_file_path, "rb") as f:
                         image_bytes = f.read()
                    
                    # Process image for OCR
                    ocr_text = process_handwritten_image(image_bytes)
                    ocr_texts_list.append(ocr_text)
                    print(f"OCR processing completed for image {i}.") # Debug print

                except Exception as e:
                     print(f"Error processing image file {i}: {str(e)}") # Debug print
                     ocr_texts_list.append(f"[Erreur de traitement image fichier {i}: {str(e)}]")

        # Process PDF Files
        if pdf_paths:
            print(f"Processing {len(pdf_paths)} PDF file(s)...") # Debug print
            for i, pdf_file_path in enumerate(pdf_paths):
                try:
                    # Read PDF bytes from saved file
                    with open(pdf_file_path, "rb") as f:
                         pdf_bytes = f.read()
                    
                    # Process PDF
                    pdf_result = process_pdf(pdf_bytes)
                    pdf_results_list.append(pdf_result)
                    print(f"PDF processing completed for file {i}.") # Debug print

                except Exception as e:
                     print(f"Error processing PDF file {i}: {str(e)}") # Debug print
                     pdf_results_list.append({"summary": f"[Erreur de traitement PDF fichier {i}: {str(e)}]", "acronyms": {}})

        # --- Processing logic ends here ---

        # --- PV Generation logic will go here ---
        print("Starting PV generation...") # Debug print

        # Call the new function to generate PV text
        generated_pv_text = await generate_pv_text_with_gemini(
            meeting_info,
            video_transcript,
            audio_transcripts_list,
            ocr_texts_list,
            pdf_results_list
        )

        # --- Return response ---
        print("PV generation process completed.") # Debug print

        # 3. Create Word document
        print("Creating Word document...") # Debug print
        word_document_buffer = create_word_pv_document(generated_pv_text, meeting_info)

        # 4. Return Word document as a StreamingResponse
        date_for_filename = meeting_info.get('date', 'N/A').replace('/', '_').replace('-', '_')
        filename = f"Procès-Verbal_{date_for_filename}.docx"

        return StreamingResponse(
            iter([word_document_buffer.getvalue()]),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f'attachment; filename="{filename}"'
            }
        )

# Uncomment to run directly
if __name__ == "__main__":
    import uvicorn
    uvicorn.run("app:app", host="0.0.0.0", port=8000, reload=True)